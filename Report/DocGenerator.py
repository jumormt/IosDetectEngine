#coding=utf-8
from docx import Document
import data
from ParseStaticXML import XMLParser

class Generator:

    def __init__(self):
        self.document = Document()

    def generate(self):
        self.write_head()
        self.write_app_info()
        self.write_binary_info()
        self.write_transport_info()
        self.write_storage_info()
        if data.static_type == 'xml':
            self.write_static_results()
        self.write_nessus_url()

        data.report_path = data.root + '/temp/{}/report/{}.doc'.format(data.start_time, data.app_bundleID)
        self.document.save(data.report_path)

        #write app info

    def write_head(self):
        self.document.add_heading(u"iOS应用漏洞检测报告", level=0)


    def write_app_info(self):
        self.document.add_heading(u"应用详情",level=1)

        self.document.add_paragraph(u'检测时间：', style='List Bullet')
        self.document.add_paragraph(data.start_time)
        self.document.add_paragraph(u'应用ID：', style='List Bullet')
        self.document.add_paragraph(data.app_bundleID)
        self.document.add_paragraph(u'应用名称：', style='List Bullet')
        self.document.add_paragraph(data.metadata['name'])
        self.document.add_paragraph(u'应用版本：', style='List Bullet')
        self.document.add_paragraph(data.metadata['app_version'])
        self.document.add_paragraph(u'应用路径：', style='List Bullet')
        self.document.add_paragraph(data.metadata['bundle_directory'])
        self.document.add_paragraph(u'沙盒路径：', style='List Bullet')
        self.document.add_paragraph(data.metadata['data_directory'])

    def write_binary_info(self):
        self.document.add_heading(u"二进制检测结果：", level=1)

        self.document.add_heading(u"二进制保护措施检测", level=2)
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        head_cell = table.rows[0].cells
        head_cell[0].text = u'架构'
        head_cell[1].text = u'加密保护(Encrypted)'
        head_cell[2].text = u'栈保护(Stack Canaries)'
        head_cell[3].text = u'自动内存管理(ARC)'
        head_cell[4].text = u'地址随机化(PIE)'
        for arch in data.protection_check_lables.keys():
            row_cells = table.add_row().cells
            row_cells[0].text = arch
            row_cells[1].text = str(data.protection_check_lables[arch]['Encrypted'])
            row_cells[2].text = str(data.protection_check_lables[arch]['Stack Canaries'])
            row_cells[3].text = str(data.protection_check_lables[arch]['ARC'])
            row_cells[4].text = str(data.protection_check_lables[arch]['PIE'])

        self.document.add_heading(u"依赖库检测(参考：{}个)".format(len(data.shared_lib)), level=2)
        table_share_lib = self.document.add_table(rows=1, cols=1)
        table_share_lib.style = 'Table Grid'
        head_cell_lib = table_share_lib.rows[0].cells
        head_cell_lib[0].text = 'Shared Library'
        for lib in data.shared_lib:
            row_lib_cells = table_share_lib.add_row().cells
            row_lib_cells[0].text = lib

        self.document.add_heading(u"硬编码(参考：{}个)".format(len(data.hardcode)), level=2)
        table_hard_code = self.document.add_table(rows=1, cols=1)
        table_hard_code.style = 'Table Grid'
        head_hard_code = table_hard_code.rows[0].cells
        head_hard_code[0].text = 'HardCode:'
        if len(data.hardcode) == 0:
            row_lib_cells = table_hard_code.add_row().cells
            row_lib_cells[0].text = u'无'
        for code in data.hardcode:
            row_lib_cells = table_hard_code.add_row().cells
            row_lib_cells[0].text = code

        url_crash_count = 0
        for key in data.fuzz_result:
            if data.fuzz_result[key]:
                url_crash_count += 1
        self.document.add_heading(u"URL Fuzz(低危：{}个)".format(url_crash_count), level=2)
        table_fuzz = self.document.add_table(rows=1, cols=2)
        table_fuzz.style = 'Table Grid'
        head_fuzz = table_fuzz.rows[0].cells
        head_fuzz[0].text = 'URL'
        head_fuzz[1].text = 'Result'
        for key in data.fuzz_result:
            row_cells = table_fuzz.add_row().cells
            row_cells[0].text = key
            if data.fuzz_result[key]:
                row_cells[1].text = 'Crash'
            else:
                row_cells[1].text = 'Safe'
        self.document.add_paragraph(u"应用程序通过appDelegate中的-[application: openURL: sourceApplicaiton: annotation:]处理接受到的URL")

        self.document.add_heading(u"反注入检测(中危)", level=2)
        if data.segment_dict.has_key('__RESTRICT') and '__restrict' in data.segment_dict['data.segment_dict']:
            self.document.add_paragraph(u"采用了[__RESTRICT,__restrict]，具备一定的反注入能力")
        else:
            self.document.add_paragraph(u"不具备反注入能力，建议添加[__RESTRICT,__restrict]")
            self.document.add_paragraph(u"攻击者可以通过注入dylib对应用中的方法和函数进行劫持和替换，严重威胁应用安全。")


    def write_transport_info(self):
        self.document.add_heading(u"传输层检测结果", level=1)
        all_item_count = 0
        for key in data.mitm_results.keys():
            all_item_count += data.mitm_results[key]
        self.document.add_heading(u"中间人攻击检测(中危：{}个)".format(all_item_count), level=2)
        if len(data.mitm_results.keys()) == 0:
            self.document.add_paragraph(u"未发现中间人攻击漏洞")
        else:
            cols_count = len(data.mitm_results.keys())+1
            table_1 = self.document.add_table(rows=1, cols=cols_count)
            table_1.style = 'Table Grid'
            head_1_cell = table_1.rows[0].cells
            head_1_cell[0].text = u'漏洞类型'
            i=1
            for item in data.mitm_results.keys():
                head_1_cell[i].text = data.mitm_results.keys()[i-1]
                i+=1
            row_1_cell = table_1.add_row().cells
            row_1_cell[0].text = u'漏洞统计'
            i=1
            for key in data.mitm_results.keys():
                row_1_cell[i].text = str(data.mitm_results[key])
                i+=1

        self.document.add_heading(u"不安全协议使用(中危：{}个)".format(len(data.traffic_unsafe_result)), level=2)
        table_2 = self.document.add_table(rows=1,cols=1)
        table_2.style = 'Table Grid'
        head_2_cell = table_2.rows[0].cells
        head_2_cell[0].text = u'未使用https'
        for item in data.traffic_unsafe_result:
            row_2_cell = table_2.add_row().cells
            row_2_cell[0].text = item
        if len(data.traffic_unsafe_result)==0:
            row_2_cell = table_2.add_row().cells
            row_2_cell[0].text = u"无"

        if "traffic" not in data.dynamic_sensitive_json:
            count = 0
        else:
            count = len(data.dynamic_sensitive_json['traffic'])
        self.document.add_heading(u"敏感信息传输(高危：{}个)".format(count), level=2)
        table_3 = self.document.add_table(rows=1, cols=3)
        table_3.style = 'Table Grid'
        head_3_cell = table_3.rows[0].cells
        head_3_cell[0].text = 'url'
        head_3_cell[1].text = 'body'
        head_3_cell[2].text = u'敏感内容'
        try:
            for item in data.dynamic_sensitive_json['traffic']:
                # print item
                row_3_cell = table_3.add_row().cells
                row_3_cell[0].text = item[0]['url']
                if item[0].has_key('body'):
                    row_3_cell[1].text = item[0]['body']
                else:
                    row_3_cell[1].text = u"无"
                row_3_cell[2].text = '-'.join(item[1])
        except Exception, e:
            pass

    def write_storage_info(self):
        self.document.add_heading(u"存储层检测结果", level=1)

        self.document.add_heading(u"动态检测结果", level=2)
        self.document.add_heading(u"KeyChain(中危：{}个)".format(len(data.dynamic_sensitive_json['keychain'])), level=3)
        table1 = self.document.add_table(rows=1,cols=3)
        table1.style = 'Table Grid'
        head1_cells = table1.rows[0].cells
        head1_cells[0].text = u'函数名称'
        head1_cells[1].text = u'函数参数'
        head1_cells[2].text = u'敏感数据'
        for item in data.dynamic_sensitive_json['keychain']:
            row1_cells = table1.add_row().cells
            row1_cells[0].text = item[0]['function']
            if item[0].has_key('attributes'):
                if item[0]['attributes'].has_key('kSecValueData'):
                    row1_cells[1].text = item[0]['attributes']['kSecValueData']
            elif item[0].has_key('attributesToUpdate'):
                if item[0]['attributesToUpdate'].has_key('kSecValueData'):
                    row1_cells[1].text = item[0]['attributesToUpdate']['kSecValueData']
            row1_cells[2].text = '-'.join(item[1])

        self.document.add_heading(u"NSUserDefaults(中危：{}个)".format(len(data.dynamic_sensitive_json['nsuserdefaults'])), level=3)
        table2 = self.document.add_table(rows=1, cols=4)
        table2.style = 'Table Grid'
        head2_cells = table2.rows[0].cells
        head2_cells[0].text = u'key'
        head2_cells[1].text = u'写入内容'
        head2_cells[2].text = u'数据类型'
        head2_cells[3].text = u'敏感数据'
        for item in data.dynamic_sensitive_json['nsuserdefaults']:
            row2_cells = table2.add_row().cells
            row2_cells[0].text = item[0]['key']
            row2_cells[1].text = str(item[0]['content'])
            row2_cells[2].text = item[0]['sourceType']
            row2_cells[3].text = '-'.join(item[1])

        self.document.add_heading(u"Plist(中危：{}个)".format(len(data.dynamic_sensitive_json['plist'])), level=3)
        table3 = self.document.add_table(rows=1, cols=4)
        table3.style = 'Table Grid'
        head3_cells = table3.rows[0].cells
        head3_cells[0].text = u'文件路径'
        head3_cells[1].text = u'写入内容'
        head3_cells[2].text = u'数据类型'
        head3_cells[3].text = u'敏感数据'
        for item in data.dynamic_sensitive_json['plist']:
            row3_cells = table3.add_row().cells
            if item[0].has_key('filepath'):
                row3_cells[0].text = item[0]['filepath']
            elif item[0].has_key('url'):
                row3_cells[0].text = item[0]['url']
            row3_cells[1].text = str(item[0]['content'])
            row3_cells[2].text = item[0]['sourceType']
            row3_cells[3].text = '-'.join(item[1])


        self.document.add_heading(u"本地审计检测结果", level=2)
        self.document.add_heading(u"KeyChain(中危：{}个)".format(len(data.keychain_values)), level=3)
        table4 = self.document.add_table(rows=1, cols=1)
        table4.style = 'Table Grid'
        head4_cells = table4.rows[0].cells
        head4_cells[0].text = u'可疑的KeyChain信息'
        if len(data.keychain_values) == 0:
            row4_cells = table4.add_row().cells
            row4_cells[0].text = u'无'
        else:
            for item in data.keychain_values:
                row4_cells = table4.add_row().cells
                row4_cells[0].text = item

        self.document.add_heading(u"数据库文件", level=3)
        if data.db_file_results:
            count = 0
            for key in data.db_file_results["input"]:
                count += len(data.db_file_results["input"][key])
            self.document.add_heading(u"用户隐私检测(中危：{}个)".format(count), level=4)
            for file_path in data.db_file_results["input"].keys():
                self.document.add_paragraph(u'文件路径：'+file_path)
                table5 = self.document.add_table(rows=1,cols=3)
                table5.style = 'Table Grid'
                head5_cells = table5.rows[0].cells
                head5_cells[0].text = u'表名'
                head5_cells[1].text = u'行信息'
                head5_cells[2].text = u'敏感数据'
                for item in data.db_file_results["input"][file_path]:
                    row5_cells = table5.add_row().cells
                    row5_cells[0].text = item[0]
                    try:
                        row5_cells[1].text = str(item[1])
                    except:
                        row5_cells[1].text = u"内容无法正常显示"
                    row5_cells[2].text = item[2]

            count = 0
            for key in data.db_file_results["keyiv"]:
                count += len(data.db_file_results["keyiv"][key])
            self.document.add_heading(u"密钥信息检测(高危：{}个)".format(count), level=4)
            for file_path in data.db_file_results["keyiv"].keys():
                self.document.add_paragraph(u'文件路径：'+file_path)
                table5 = self.document.add_table(rows=1,cols=3)
                table5.style = 'Table Grid'
                head5_cells = table5.rows[0].cells
                head5_cells[0].text = u'表名'
                head5_cells[1].text = u'行信息'
                head5_cells[2].text = u'敏感数据'
                for item in data.db_file_results["keyiv"][file_path]:
                    row5_cells = table5.add_row().cells
                    row5_cells[0].text = item[0]
                    try:
                        row5_cells[1].text = str(item[1])
                    except:
                        row5_cells[1].text = u"内容无法正常显示"
                    row5_cells[2].text = item[2]
            count = 0
            for key in data.db_file_results["txt"]:
                count += len(data.db_file_results["txt"][key])
            self.document.add_heading(u"自定义项检测(高危：{}个)".format(count), level=4)
            for file_path in data.db_file_results["txt"].keys():
                self.document.add_paragraph(u'文件路径：'+file_path)
                table5 = self.document.add_table(rows=1, cols=3)
                table5.style = 'Table Grid'
                head5_cells = table5.rows[0].cells
                head5_cells[0].text = u'表名'
                head5_cells[1].text = u'行信息'
                head5_cells[2].text = u'敏感数据'
                for item in data.db_file_results["txt"][file_path]:
                    row5_cells = table5.add_row().cells
                    row5_cells[0].text = item[0]
                    try:
                        row5_cells[1].text = str(item[1])
                    except:
                        row5_cells[1].text = u"内容无法正常显示"
                    row5_cells[2].text = item[2]
        else:
            self.document.add_paragraph(u"无数据库文件")

        self.document.add_heading(u"Plist文件", level=3)
        if data.plist_file_results:
            count = 0
            for key in data.plist_file_results["input"]:
                count += len(data.plist_file_results["input"][key])
            self.document.add_heading(u"用户隐私检测(中危：{}个)".format(count), level=4)
            for file_path in data.plist_file_results["input"].keys():
                self.document.add_paragraph(u'文件路径：'+file_path)
                table6 = self.document.add_table(rows=1,cols=3)
                table6.style = 'Table Grid'
                head6_cells = table6.rows[0].cells
                head6_cells[0].text = u'key路径'
                head6_cells[1].text = u'单元信息'
                head6_cells[2].text = u'敏感数据'
                for item in data.plist_file_results["input"][file_path]:
                    row6_cells = table6.add_row().cells
                    row6_cells[0].text = item[0]
                    row6_cells[1].text = str(item[1])
                    row6_cells[2].text = item[2]

            count = 0
            for key in data.plist_file_results["keyiv"]:
                count += len(data.plist_file_results["keyiv"][key])
            self.document.add_heading(u"密钥信息检测(高危：{}个)".format(count), level=4)
            for file_path in data.plist_file_results["keyiv"].keys():
                self.document.add_paragraph(u'文件路径：'+file_path)
                table6 = self.document.add_table(rows=1,cols=3)
                table6.style = 'Table Grid'
                head6_cells = table6.rows[0].cells
                head6_cells[0].text = u'key路径'
                head6_cells[1].text = u'单元信息'
                head6_cells[2].text = u'敏感数据'
                for item in data.plist_file_results["keyiv"][file_path]:
                    row6_cells = table6.add_row().cells
                    row6_cells[0].text = item[0]
                    row6_cells[1].text = str(item[1])
                    row6_cells[2].text = item[2]

            count = 0
            for key in data.plist_file_results["txt"]:
                count += len(data.plist_file_results["txt"][key])
            self.document.add_heading(u"自定义项检测(高危：{}个)".format(count), level=4)
            for file_path in data.plist_file_results["txt"].keys():
                self.document.add_paragraph(u'文件路径：'+file_path)
                table6 = self.document.add_table(rows=1,cols=3)
                table6.style = 'Table Grid'
                head6_cells = table6.rows[0].cells
                head6_cells[0].text = u'key路径'
                head6_cells[1].text = u'单元信息'
                head6_cells[2].text = u'敏感数据'
                for item in data.plist_file_results["txt"][file_path]:
                    row6_cells = table6.add_row().cells
                    row6_cells[0].text = item[0]
                    row6_cells[1].text = str(item[1])
                    row6_cells[2].text = item[2]
        else:
            self.document.add_paragraph(u"无plist文件")

        self.document.add_heading(u"Log审计", level=3)
        if data.log_file_results:
            count = 0
            for key in data.log_file_results["input"]:
                count += len(data.log_file_results["input"][key])
            self.document.add_heading(u"用户隐私检测(中危：{}个)".format(count), level=4)
            for file_path in data.plist_file_results["input"].keys():
                self.document.add_paragraph(u'文件路径：' + file_path)
                table6 = self.document.add_table(rows=1, cols=2)
                table6.style = 'Table Grid'
                head6_cells = table6.rows[0].cells
                head6_cells[0].text = u'log信息'
                head6_cells[1].text = u'敏感信息'
                for item in data.plist_file_results["input"][file_path]:
                    row6_cells = table6.add_row().cells
                    row6_cells[0].text = item[0]
                    row6_cells[2].text = item[1]

            count = 0
            for key in data.log_file_results["keyiv"]:
                count += len(data.log_file_results["keyiv"][key])
            self.document.add_heading(u"密钥信息检测(高危：{}个)".format(count), level=4)
            for file_path in data.plist_file_results["keyiv"].keys():
                self.document.add_paragraph(u'文件路径：' + file_path)
                table6 = self.document.add_table(rows=1, cols=2)
                table6.style = 'Table Grid'
                head6_cells = table6.rows[0].cells
                head6_cells[0].text = u'log信息'
                head6_cells[1].text = u'敏感信息'
                for item in data.plist_file_results["keyiv"][file_path]:
                    row6_cells = table6.add_row().cells
                    row6_cells[0].text = item[0]
                    row6_cells[2].text = item[1]

            count = 0
            for key in data.log_file_results["txt"]:
                count += len(data.log_file_results["txt"][key])
            self.document.add_heading(u"自定义项检测(高危：{}个)".format(count), level=4)
            for file_path in data.plist_file_results["txt"].keys():
                self.document.add_paragraph(u'文件路径：' + file_path)
                table6 = self.document.add_table(rows=1, cols=2)
                table6.style = 'Table Grid'
                head6_cells = table6.rows[0].cells
                head6_cells[0].text = u'log信息'
                head6_cells[1].text = u'敏感信息'
                for item in data.plist_file_results["txt"][file_path]:
                    row6_cells = table6.add_row().cells
                    row6_cells[0].text = item[0]
                    row6_cells[2].text = item[1]
        else:
            self.document.add_paragraph(u"无敏感信息")

        self.document.add_heading(u"本地文件保护检测(参考：{}个)".format(len(data.local_file_protection)), level=3)
        table7 = self.document.add_table(rows=1,cols=2)
        table7.style = 'Table Grid'
        head7_cells = table7.rows[0].cells
        head7_cells[0].text = u'文件路径'
        head7_cells[1].text = u'保护信息'
        for item in data.local_file_protection:
            row7_cells = table7.add_row().cells
            row7_cells[0].text = item[0]
            protection_info = item[1].strip()
            if protection_info == u"NSFileProtectionNone":
                row7_cells[1].text = u"无保护" + protection_info
            else:
                row7_cells[1].text = u"安全" + protection_info


    def write_static_results(self):
        xml_parser = XMLParser()
        xml_parser.start_parse()
        unsafe_matchs = xml_parser.unsecure_match_results
        safes_unmatch = xml_parser.secure_not_match_results
        # print unsafe_matchs
        # print safes_unmatch
        self.document.add_heading(u"反汇编检测结果", level=1)
        self.document.add_heading(u"不安全代码检测结果", level=2)
        if len(unsafe_matchs.keys())==0:
            self.document.add_paragraph(u"无内容")
        for fun in unsafe_matchs.keys():
            self.document.add_heading(fun, level=3)
            table = self.document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            head_cell = table.rows[0].cells
            head_cell[0].text = u'规则id'
            head_cell[1].text = u'规则名称'
            head_cell[2].text = u'规则描述'
            head_cell[3].text = u'解决方法'
            head_cell[4].text = u'危险等级'
            for rule in unsafe_matchs[fun]:
                cells = table.add_row().cells
                cells[0].text = rule._id
                cells[1].text = rule._name
                cells[2].text = rule._desc
                cells[3].text = rule._solution
                if rule._risk_level.strip() == u'1':
                    cells[4].text = u"低"
                elif rule._risk_level.strip() == u'2':
                    cells[4].text = u"中"
                elif rule._risk_level.strip() == u'3':
                    cells[4].text = u"高"
                else:
                    cells[4].text = rule._risk_level

        self.document.add_heading(u"保护措施检测结果", level=2)
        if len(safes_unmatch)==0:
            self.document.add_paragraph(u"无内容")
        else:
            table = self.document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            head_cell = table.rows[0].cells
            head_cell[0].text = u'规则id'
            head_cell[1].text = u'规则名称'
            head_cell[2].text = u'规则描述'
            head_cell[3].text = u'解决方法'
            head_cell[4].text = u'危险等级'
            for rule in safes_unmatch:
                cells = table.add_row().cells
                cells[0].text = rule._id
                cells[1].text = rule._name
                cells[2].text = rule._desc
                cells[3].text = rule._solution
                # cells[4].text = rule._risk_level
                if rule._risk_level.strip() == u'1':
                    cells[4].text = u"低"
                elif rule._risk_level.strip() == u'2':
                    cells[4].text = u"中"
                elif rule._risk_level.strip() == u'3':
                    cells[4].text = u"高"
                else:
                    cells[4].text = rule._risk_level


    def write_nessus_url(self):
        self.document.add_heading(u"服务器扫描结果", level=1)
        self.document.add_paragraph(data.nessus_url)